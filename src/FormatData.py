###############################
##--YEARBOOK DATA FORMATTER--##
#######--HASHIM IQBAL--########
#######--11/01/2021--##########
###############################

#READ DOCUMENTATION FILE BEFORE EXECUTING APPLICATION

import tkinter as tk
from tkinter import filedialog, Text
import time
import os
import docx
import pandas as pd

#Global variable which will open up the file in other functions 
globalPath = ''

#Function to add a file's directory to the frame
def addFile():

    #This tells the function that appPath is global in the context of the function
    #Without this, a new local variable of the same name is created, which we don't want.
    global globalPath

    #If there are any widgets in the frame (filenames), destroy them before putting in the updated one.
    for widget in files.winfo_children():
        widget.destroy()

    #Opens the file dialog to allow users to select anything of filetype xlsx.
    filePath = filedialog.askopenfilename(initialdir="/", title='Select File', 
    filetypes=(('Excel Spreadsheets', '*.xlsx'), ('All Files', '*.*')))

    #Truncates the variable to just have one path and not multiple.
    #Stores the path in our global variable, which will be accessed in a later function
    globalPath = ''
    globalPath += filePath

    #Some validation. If the path is empty, don't bother making a label for it.
    #Otherwise, make a label.
    if filePath == '':
        pass
    else:
        pathLabel = tk.Label(files, text=filePath)
        pathLabel.pack()


#Reads the .xlsx file using pandas, and writes to a .docx file in the desired format.
def readAndconvert():
    #This bit of code only exists to time how long the function takes to run.
    t1 = time.time()

    #We need to declare this in functions that actively use our global path.
    global globalPath

    #If the frame doesn't have any path labels, write an error label.
    #Otherwise, do the full read and convert to .docx
    if len(files.winfo_children()) < 1:
        errorMessage = tk.Label(files, text='Error. No file selected', fg='#FF0000')
        errorMessage.pack()
    else:
        #Using pandas to read the data from the .xlsx file
        #Then rename the columns so that they can be indexed more easily.
        #Drop any duplicate records and only keep the most recent ones.
        df = pd.read_excel(globalPath, usecols='B:N', index_col=0)
        df.columns = ['Name', 'Nickname', 'Year Joined', 'Memorable Moments', 'Friends', 'Wannabe', 'Quotes', 'Embarrassing Moments', 'In 10 Years...', 'Paragraph', 'Confirmation', 'Writers']
        df = df.drop_duplicates('Name', keep='last')

        #Create a 2D array to hold each profile
        profiles = [[None]*12 for i in range(len(df))]
        
        #Iterate through the dataframe and add the records to the array.
        for x in range(len(profiles)):
            for y in range(len(profiles[x])):
                profiles[x][y] = df.iloc[x, y]

        #Create a document using the python-docx module
        document = docx.Document()

        #Configure the font the text will be in. 
        #This doesn't really matter, but Arial is something that is easier to read.
        font = document.styles['Normal'].font
        font.name = 'Arial'

        #Iterate through each profile, and write the content to the file.
        for x in range(len(profiles)):
            r = document.add_paragraph()
            r.paragraph_format.space_after = docx.shared.Pt(0.5)
            runner = r.add_run(str(profiles[x][0]))
            runner.bold = True

            p = document.add_paragraph(
                                        str(profiles[x][1]) + '\n'
                                        + str(profiles[x][2]) + '\n'
                                        + str(profiles[x][3]) + '\n'
                                        + str(profiles[x][4]) + '\n'
                                        + str(profiles[x][5]) + '\n'
                                        + str(profiles[x][6]) + '\n'
                                        + str(profiles[x][7]) + '\n'
                                        + str(profiles[x][8]) + '\n'
                                        + str(profiles[x][9]) + '\n'
                                        + str(profiles[x][11])
                                        + '\n'
                                    )

        #Use a filedialog box to ask the user where they wish to save the converted data and as what file name.
        save_location = filedialog.asksaveasfilename(defaultextension='.docx', initialdir='/', title='Save File', filetypes=(('Word Documents', '*.docx'), ('All Files', '*.*')))
        #If save_location has been initialised, save the document. Otherwise, don't do anything.
        if save_location:
            #Save the document (no other operations with the doc can be carried out after this.)
            #Document will be saved in the root directory of the folder.
            document.save(save_location)
        else:
            pass
    
    t2 = time.time()
    print('The time taken by the conversion algorithm is ',t2-t1, ' seconds')



#Create a GUI window
master = tk.Tk()

#Give this window a title and icon
master.title('Yearbook Data Formatter')
master.iconbitmap('./src/PyIcon.ico')

#Disables resizing buttons on mac and windows
#Also gives the window a fixed size of 700x700 px
master.resizable(0, 0)
master.geometry('520x600')

#Canvases are better for displaying elements than windows
canvas = tk.Canvas(master, height=500, width=500, bg='#263D42')
canvas.pack()

#Creates a frame which has a different colour and provides contrast
files = tk.Frame(master, bg='white')
files.place(relwidth=0.8, relheight=0.5, relx=0.10, rely=0.25)

#Button which will call the addFile function when clicked
openFile = tk.Button(master, text='Open File', padx=10, pady=5, fg='white', bg='#263D42', command=addFile)
openFile.pack()

#Button which will call the readAndconvert function when clicked
convertData = tk.Button(master, text='Convert Data', padx=10, pady=5, fg='white', bg='#263D42', command=readAndconvert)
convertData.pack()

exitButton = tk.Button(master, text='Exit Program', padx=10, pady=5, fg='white', bg='#263D42', command=master.quit)
exitButton.pack()

#This line basically tells tkinter that we're done with the GUI window, and no further changes will be made.
master.mainloop()